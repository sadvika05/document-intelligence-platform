# backend/api/views.py
import os
import uuid
import logging
import fitz  # PyMuPDF for PDF
from docx import Document as DocxDocument  # For DOCX
from rest_framework.decorators import api_view, parser_classes
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from .models import Document, DocumentChunk
from .serializers import DocumentSerializer
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings
import openai
import requests
from dotenv import load_dotenv

logger = logging.getLogger(__name__)

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
USE_OPENAI = openai.api_key is not None
LM_STUDIO_URL = "http://localhost:1234/v1/chat/completions"

model = SentenceTransformer("all-MiniLM-L6-v2")
client = chromadb.PersistentClient(path="./chroma_data")
collection = client.get_or_create_collection("documents")

def extract_text_and_pages(file_path, content_type):
    if content_type == "text/plain":
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read(), 1
    elif content_type == "application/pdf":
        doc = fitz.open(file_path)
        text = "\n\n".join([page.get_text() for page in doc])
        return text, len(doc)
    elif content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        doc = DocxDocument(file_path)
        text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text, len(doc.paragraphs) // 30 or 1
    else:
        raise ValueError("Unsupported file type")

@api_view(['POST'])
@parser_classes([MultiPartParser, FormParser])
def upload_document(request):
    try:
        if 'file' not in request.FILES:
            return Response({"status": "error", "detail": "No file uploaded"}, status=400)

        file = request.FILES['file']
        upload_dir = 'uploads'
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, f"{uuid.uuid4()}_{file.name}")

        with open(file_path, 'wb+') as f:
            for chunk in file.chunks():
                f.write(chunk)

        try:
            text, pages = extract_text_and_pages(file_path, file.content_type)
        except ValueError as ve:
            return Response({"status": "error", "detail": str(ve)}, status=400)

        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        document = Document.objects.create(
            title=file.name,
            file_path=file_path,
            file_type=file.content_type,
            size=file.size,
            pages=pages,
            status="processed"
        )

        for i, chunk in enumerate(paragraphs):
            emb = model.encode(chunk).tolist()
            embedding_id = str(uuid.uuid4())
            collection.add(documents=[chunk], embeddings=[emb], metadatas=[{"source": str(document.id), "chunk": i}], ids=[embedding_id])
            DocumentChunk.objects.create(document=document, chunk_index=i, text=chunk, embedding_id=embedding_id)

        return Response({"status": "success", "document_id": document.id})
    except Exception as e:
        logger.error(f"Upload error: {e}")
        return Response({"status": "error", "detail": str(e)}, status=500)

@api_view(['POST'])
def ask_question(request):
    import time
    start_time = time.time()

    doc_id = request.data.get("document_id")
    question = request.data.get("question")
    top_k = int(request.data.get("top_k", 3))

    try:
        doc = Document.objects.get(id=doc_id)
        if doc.status != 'processed':
            return Response({
                "success": False,
                "error": f"Document not ready. Status: {doc.status}"
            }, status=400)
    except Document.DoesNotExist:
        return Response({"success": False, "error": "Document not found"}, status=404)

    try:
        question_embedding = model.encode(question).tolist()
        results = collection.query(
            query_embeddings=[question_embedding],
            n_results=top_k,
            where={"source": str(doc_id)}
        )

        if not results['documents'][0]:
            return Response({
                "success": False,
                "error": "No relevant content found for your question"
            }, status=404)

        context_chunks = results['documents'][0]
        context = "\n".join(context_chunks)
        prompt = f"""You are a helpful assistant. Use the provided document context to answer the user's question.

Context:
{context}

Question: {question}

Answer:"""

        if USE_OPENAI:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant for answering document-based questions."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=500
            )
            answer = response['choices'][0]['message']['content'].strip()
        else:
            response = requests.post(LM_STUDIO_URL, json={
                "model": "TheBloke/Mistral-7B-Instruct-v0.1-GGUF",
                "messages": [
                    {"role": "system", "content": "You are a helpful assistant for answering document-based questions."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 500
            })
            answer = response.json()["choices"][0]["message"]["content"].strip()

        return Response({
            "success": True,
            "question": question,
            "answer": answer,
            "sources": results['metadatas'][0],
            "chunks_used": len(context_chunks),
            "processing_time": round(time.time() - start_time, 2)
        })

    except Exception as e:
        logger.error(f"Query failed: {e}")
        return Response({
            "success": False,
            "error": "Failed to generate answer",
            "details": str(e),
            "processing_time": round(time.time() - start_time, 2)
        }, status=500)

@api_view(['GET'])
def list_documents(request):
    documents = Document.objects.all()
    serializer = DocumentSerializer(documents, many=True)
    return Response(serializer.data)
